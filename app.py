import os
import shutil
import subprocess
import time
from pathlib import Path
from docx import Document
from langchain_community.document_loaders import TextLoader, PyPDFLoader, UnstructuredHTMLLoader
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter 
from langchain.chains import RetrievalQA
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.docstore.document import Document as LcDocument
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import re 
import json 

# --- ANSI Color Codes ---
# Define these as raw escape sequences, not Python variable names for LLM prompt
ANSI_YELLOW = "\033[1;33m" # Bright Yellow
ANSI_BLUE = "\033[1;34m"   # Bright Blue
ANSI_RESET = "\033[0m"    # Reset to default terminal color
ANSI_ORANGE = "\033[38;5;208m" # Corrected ANSI Orange code
ANSI_LIGHT_BLUE = "\033[94m"   # Light Blue

# --- Configuration ---
# Base directory for all your document data
BASE_DATA_DIR = Path("data")

# Specific subdirectories for different document types
MISC_DIR = BASE_DATA_DIR / "misc" 
SEC_DIR = BASE_DATA_DIR / "sec"
LEGAL_DIR = BASE_DATA_DIR / "legal"
NEWS_DIR = BASE_DATA_DIR / "news" # Directory for raw news articles (text + json)

# NEW: Directories for blogs and social media content
BLOG_DIR = BASE_DATA_DIR / "blogs" 
SOCIAL_MEDIA_DIR = BASE_DATA_DIR / "social_media"

# Base directory for all vector stores (each will get its own subdirectory)
BASE_VECTOR_DB_DIR = Path("vectorstores") 

LLAMAFILE_NAME = "Mistral-7B-Instruct-v0.3.Q5_K_M.llamafile" 
LLAMAFILE_API_BASE = "http://localhost:8080/v1"
GPU_OFFLOAD_LAYERS = 999 

# --- Logging Configuration ---
LLAMAFILE_LOG_FILE = "llamafile_server.log"

# --- LLM Prompts (Defined globally as they are constant) ---

# NEW: Headline Summarizer Prompt for direct listing
# This prompt will now be used ONLY when the LLM is involved in synthesizing summaries.
# For direct file system reads, the app will format the output directly.
HEADLINE_SUMMARIZER_PROMPT = f"""You are a helpful assistant.
Given a list of news article titles and their brief summaries, your task is to present them clearly.
Do NOT attempt to group them into themes or categories. Simply list each article entry as provided.
Format each article entry exactly as: "{ANSI_YELLOW}Title:{ANSI_RESET} {ANSI_ORANGE}[Article Title]{ANSI_RESET} {ANSI_YELLOW}Summary:{ANSI_RESET} {ANSI_LIGHT_BLUE}[Article Summary]{ANSI_RESET}"
Make sure to include the ANSI color codes as specified.

Actual Input:
{{context}}

Output:"""

# System Instruction Prompt for general Q&A
SYSTEM_INSTRUCTION_PROMPT = """You are a helpful, accurate, and concise assistant.
Prioritize answering questions based on the provided document context.
If information for the question is found in the document, preface your answer or relevant section with: "**Based on the provided document [Source: source_filename, Title: original_title if available]:**"
If the provided context is insufficient or the question requires broader knowledge, preface your answer or relevant section with: "**Drawing on general knowledge:**"
If the answer cannot be found in the provided document, explicitly state: "Information for this question is not available in the provided document." before potentially offering a general answer if appropriate."""


# --- Helper Functions ---

def is_headline_query(query):
    """
    Determines if a query is likely a request for general news headlines/overview,
    or general summaries (blogs, social media).
    More flexible matching.
    """
    query_lower = query.lower()
    headline_keywords = [
        "summarize headlines", "latest news", "news overview", "recent news",
        "daily briefing", "what's the news", "tell me about the news", "whats happening",
        "headlines", "news", # News specific
        "summarize blogs", "blog updates", "blog summaries", "blog posts", # Blog specific
        "summarize social media", "social media trends", "what's trending", "social media overview", # Social media specific
        "get headlines", "list headlines" # Explicit keywords for direct retrieval
    ]
    # If any of these keywords are present, and it's not explicitly asking for "details" or "full article"
    if any(keyword in query_lower for keyword in headline_keywords):
        if not any(detail_kw in query_lower for detail_kw in ["details", "full article", "in depth", "explain", "describe"]):
            return True
    return False

def parse_date_and_site_from_query(query):
    """
    Parses date and site from a query string and returns a ChromaDB filter dictionary
    and the cleaned query.
    Date formats supported: "today", "yesterday", "last week", "this month", "YYYY-MM-DD".
    Site formats supported: "cnn.com", "foxnews.com", etc., and also common names like "fox news".
    """
    query_lower = query.lower()
    filters = {}
    cleaned_query = query # Initialize cleaned_query with the original query

    # --- Date Parsing ---
    today = datetime.now()
    date_filter_applied = False

    # "today"
    if "today" in query_lower:
        start_date = today.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = today.replace(hour=23, minute=59, second=59, microsecond=999999)
        filters["publication_date"] = {"$gte": start_date.isoformat(), "$lte": end_date.isoformat()}
        cleaned_query = cleaned_query.replace("today", "").strip()
        date_filter_applied = True
    # "yesterday"
    elif "yesterday" in query_lower:
        yesterday = today - timedelta(days=1)
        start_date = yesterday.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = yesterday.replace(hour=23, minute=59, second=59, microsecond=999999)
        filters["publication_date"] = {"$gte": start_date.isoformat(), "$lte": end_date.isoformat()}
        cleaned_query = cleaned_query.replace("yesterday", "").strip()
        date_filter_applied = True
    # "last week" (past 7 days including today)
    elif "last week" in query_lower:
        last_week_start = today - timedelta(days=7)
        filters["publication_date"] = {"$gte": last_week_start.isoformat()}
        cleaned_query = cleaned_query.replace("last week", "").strip()
        date_filter_applied = True
    # "this month" (start of current month)
    elif "this month" in query_lower:
        this_month_start = today.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        filters["publication_date"] = {"$gte": this_month_start.isoformat()}
        cleaned_query = cleaned_query.replace("this month", "").strip()
        date_filter_applied = True
    # Specific date YYYY-MM-DD or MM-DD-YYYY
    else:
        date_match = re.search(r'\b(\d{4}-\d{2}-\d{2}|\d{2}-\d{2}-\d{4})\b', query_lower)
        if date_match:
            date_str = date_match.group(0)
            try:
                if len(date_str.split('-')[0]) == 4: # YYYY-MM-DD
                    specific_date = datetime.strptime(date_str, "%Y-%m-%d")
                else: # MM-DD-YYYY
                    specific_date = datetime.strptime(date_str, "%m-%d-%Y")
                
                start_date = specific_date.replace(hour=0, minute=0, second=0, microsecond=0)
                end_date = specific_date.replace(hour=23, minute=59, second=59, microsecond=999999)
                filters["publication_date"] = {"$gte": start_date.isoformat(), "$lte": end_date.isoformat()}
                cleaned_query = re.sub(pattern, '', cleaned_query, flags=re.IGNORECASE).strip()
                date_filter_applied = True
            except ValueError:
                pass # Invalid date format, continue without date filter

    # --- Site Parsing ---
    # Extended regex to catch common names or full domains. Added general site keywords.
    site_patterns = {
        r'\bfox\s*news(?:\.com)?\b': 'foxnews.com',
        r'\bcnn(?:\.com)?\b': 'cnn.com',
        r'\brt(?:\.com)?\b': 'rt.com',
        r'\bpatriots\.win\b': 'patriots.win',
        r'\b4chan(?:\.org/pol)?\b': 'boards.4chan.org', # Catches "4chan" or "4chan.org/pol"
        r'\bgoogle(?:\.com/finance)?\b': 'google.com/finance',
        r'\bthemoscowtimes(?:\.com)?\b': 'themoscowtimes.com',
        r'\btass(?:\.com)?\b': 'tass.com',
        r'\beuronews(?:\.com)?\b': 'euronews.com',
        r'\baljazeera(?:\.com)?\b': 'aljazeera.com',
        r'\becns(?:\.cn)?\b': 'ecns.cn',
        # Add common blog/social media sites here if you're consistently scraping them with site_origin metadata
        r'\bmedium(?:\.com)?\b': 'medium.com',
        r'\breddit(?:\.com)?\b': 'reddit.com',
        r'\btwitter(?:\.com)?\b': 'twitter.com', # Assuming you'd get these from social media scrapes
        r'\bx\.com(?:\/.*)?\b': 'x.com' # New Twitter domain
    }

    site_match_found = False
    for pattern, domain in site_patterns.items():
        match = re.search(pattern, query_lower)
        if match:
            filters["site_origin"] = domain
            cleaned_query = cleaned_query.replace(match.group(0), "").strip()
            site_match_found = True
            break # Found a match, no need to check other patterns

    # Clean up any residual "from" or "on" prepositions if they are at the start/end of the cleaned query
    cleaned_query = re.sub(r'^(from|on)\s+|\s+(from|on)$', '', cleaned_query, flags=re.IGNORECASE).strip()

    # Remove extra spaces caused by replacements
    cleaned_query = re.sub(r'\s+', ' ', cleaned_query).strip()

    return filters, cleaned_query if filters or date_filter_applied else None # Return None if no filters were applied


def extract_text_from_docx(docx_path):
    """
    Extracts text from a .docx file.
    """
    try:
        doc = Document(docx_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX file {docx_path}: {e}")
        return ""

def load_documents_from_directory(directory: Path, load_summaries=False, summary_kb_name=None):
    """
    Loads documents from a given directory into a list of Langchain Document objects.
    Handles different directory structures based on whether it's the NEWS_DIR or others.
    """
    documents = []
    print(f"  Scanning: {directory} {'(for summaries)' if load_summaries else ''}")
    if not directory.exists():
        print(f"  Warning: Directory '{directory}' does not exist. Skipping.")
        return []

    # Determine the starting point for os.walk based on the directory type
    dirs_to_walk = []
    if directory == NEWS_DIR:
        # Check for timestamped subdirectories first
        potential_subdirs = [d for d in directory.iterdir() if d.is_dir()] 
        if potential_subdirs:
            dirs_to_walk.extend(potential_subdirs)
        else:
            # If no subdirs, treat NEWS_DIR itself as the starting point for walk
            print(f"  Info: No timestamped subdirectories found in {directory}. Scanning {directory} directly for news.")
            dirs_to_walk = [directory]
    else: # For MISC, SEC, LEGAL, BLOG, SOCIAL_MEDIA, always start walk from the main directory
        dirs_to_walk = [directory]

    # If no directories to walk (e.g., NEWS_DIR is empty and has no subdirs), return early
    if not dirs_to_walk:
        print(f"  Info: No valid content directories found within {directory}.")
        return []

    for start_dir in dirs_to_walk:
        for root, _, files in os.walk(start_dir):
            for file_name in files:
                file_path = os.path.join(root, file_name) # Define file_path here
                
                # Logic for .txt files
                if file_name.endswith(".txt"):
                    metadata_file_path = Path(file_path).with_suffix('.json')
                    
                    if load_summaries or directory in [NEWS_DIR, BLOG_DIR, SOCIAL_MEDIA_DIR]: 
                        if metadata_file_path.exists():
                            try:
                                with open(metadata_file_path, 'r', encoding='utf-8') as f:
                                    metadata = json.load(f)
                                    
                                    if "publication_date" in metadata and isinstance(metadata["publication_date"], str):
                                        try:
                                            parsed_dt = datetime.strptime(metadata["publication_date"], "%Y-%m-%d_%H-%M-%S")
                                            metadata["publication_date"] = parsed_dt.isoformat() 
                                        except ValueError:
                                            print(f"    Warning: Could not parse date format in metadata for {file_name}. Storing as original string.")
                                    
                                    content_to_load = ""
                                    if load_summaries and "article_summary" in metadata:
                                        content_to_load = metadata["article_summary"]
                                    elif not load_summaries: 
                                        loader = TextLoader(file_path, encoding='utf-8')
                                        docs_from_loader = loader.load()
                                        if docs_from_loader:
                                            content_to_load = docs_from_loader[0].page_content
                                    
                                    if content_to_load:
                                        doc_metadata = {
                                            "source": file_name, 
                                            "file_path": file_path, 
                                            "category": summary_kb_name if load_summaries and summary_kb_name else directory.name
                                        }
                                        if "site_origin" in metadata and metadata["site_origin"].startswith("www."):
                                            metadata["site_origin"] = metadata["site_origin"][4:] 
                                        doc_metadata.update(metadata) 
                                        
                                        print(f"    DEBUG: Loaded doc: {file_name}, Site: {doc_metadata.get('site_origin')}, Summary Length: {len(content_to_load)}")

                                        documents.append(LcDocument(page_content=content_to_load, metadata=doc_metadata))
                                    else:
                                        print(f"    Skipping {file_name}: No content (or summary) to load from JSON/text.")
                            except json.JSONDecodeError as e:
                                print(f"    Error reading JSON metadata for {file_name}: {e}. Skipping metadata and text loading.")
                            except Exception as e:
                                print(f"    Error processing {file_name}: {e}.")
                        else: 
                            try:
                                loader = TextLoader(file_path, encoding='utf-8')
                                docs_from_loader = loader.load()
                                for doc in docs_from_loader:
                                    doc.metadata["category"] = directory.name
                                    doc.metadata["source"] = file_name 
                                    doc.metadata["file_path"] = file_path 
                                    documents.append(doc)
                            except Exception as e:
                                print(f"    Error loading text file {file_name}: {e}")
                    else: 
                        try:
                            loader = TextLoader(file_path, encoding='utf-8')
                            docs_from_loader = loader.load()
                            for doc in docs_from_loader:
                                doc.metadata["category"] = directory.name
                                doc.metadata["source"] = file_name 
                                doc.metadata["file_path"] = file_path 
                                documents.append(doc)
                        except Exception as e:
                            print(f"    Error loading text file {file_name}: {e}")
                
                # Logic for non-TXT files (DOCX, PDF, HTML) - only loaded if not in summary mode
                elif not load_summaries and (file_name.endswith(".docx") or file_name.endswith(".pdf") or file_name.endswith(".html") or file_name.endswith(".htm")): 
                    try:
                        if file_name.endswith(".docx"):
                            print(f"    Loading DOCX file: {file_name}")
                            text_content = extract_text_from_docx(file_path) 
                            if text_content:
                                documents.append(LcDocument(page_content=text_content, metadata={"source": file_name, "file_path": file_path, "category": directory.name}))
                        elif file_name.endswith(".pdf"):
                            print(f"    Loading PDF file: {file_name}")
                            try:
                                loader = PyPDFLoader(file_path) 
                                docs = loader.load()
                                for doc in docs:
                                    doc.metadata["category"] = directory.name
                                    doc.metadata["source"] = file_name 
                                    doc.metadata["file_path"] = file_path 
                                    documents.append(doc)
                            except ImportError:
                                print(f"    Skipping PDF file {file_name}: 'pypdf' not installed. Please install with `pip install pypdf`.")
                            except Exception as pdf_e:
                                print(f"    Error loading PDF file {file_name}: {pdf_e}. Ensure 'pypdf' is installed and the PDF is not corrupted.") 
                        elif file_name.endswith(".html") or file_name.endswith(".htm"):
                            print(f"    Loading HTML file: {file_name}")
                            if directory == SEC_DIR:
                                try:
                                    print(f"      Using BeautifulSoup for enhanced SEC HTML parsing for {file_name}")
                                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                        html_content = f.read()
                                    soup = BeautifulSoup(html_content, 'html.parser')
                                    for script_or_style in soup(["script", "style"]):
                                        script_or_style.extract()
                                    text = soup.get_text()
                                    text = os.linesep.join([s for s in text.splitlines() if s.strip()])
                                    text = text.strip()
                                    if text:
                                        documents.append(LcDocument(page_content=text, metadata={"source": file_name, "file_path": file_path, "category": directory.name}))
                                    else:
                                        print(f"      No meaningful text extracted from {file_name} using BeautifulSoup. File might be malformed or empty.")
                                except ImportError:
                                    print(f"    Skipping HTML file {file_name} (SEC): 'beautifulsoup4' not installed. Please install with `pip install beautifulsoup4`.")
                                except Exception as bs_e:
                                    print(f"    Error parsing HTML file {file_name} (SEC) with BeautifulSoup: {bs_e}.")
                            else:
                                try:
                                    loader = UnstructuredHTMLLoader(file_path) 
                                    docs = loader.load()
                                    for doc in docs:
                                        doc.metadata["category"] = directory.name
                                        doc.metadata["source"] = file_name 
                                        doc.metadata["file_path"] = file_path 
                                        documents.extend(docs) 
                                except ImportError:
                                    print(f"    Skipping HTML file {file_name}: 'unstructured' or its dependencies not installed. Please install with `pip install unstructured`.")
                                except Exception as unstr_e:
                                    print(f"    Error loading HTML file {file_name} with UnstructuredHTMLLoader: {unstr_e}.")
                    except Exception as e: 
                        print(f"    An unexpected error occurred loading {file_name} as {file_name.split('.')[-1].upper()} file: {e}")
                # IMPORTANT: No 'else' here to specifically skip JSON files, as they are only processed as metadata
                # for .txt files and should not be treated as standalone documents.

    print(f"  Loaded {len(documents)} documents from {directory}.")
    return documents


def initialize_vector_store_for_category(documents, category_name: str, use_summaries=False):
    """
    Initializes or loads a ChromaDB vector store for a specific category.
    Each category gets its own persistent directory within BASE_VECTOR_DB_DIR.
    If use_summaries is True, it expects documents loaded with summaries as page_content.
    """
    category_db_path = BASE_DATA_DIR / "vectorstores" / category_name
    print(f"Initializing embeddings model for {category_name}...")
    embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")

    if documents:
        print(f"Creating/updating vector store for '{category_name}' at {category_db_path} with {len(documents)} documents...")
        if category_db_path.exists():
            shutil.rmtree(category_db_path)
            print(f"Cleared existing vector store at {category_db_path}")

        # Text splitting for full documents, but not for already short summaries
        if not use_summaries:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            splits = text_splitter.split_documents(documents)
        else:
            splits = documents # Summaries are already short enough, no need to split

        vectorstore = Chroma.from_documents(
            documents=splits,
            embedding=embeddings,
            persist_directory=str(category_db_path)
        )
        print(f"Vector store for '{category_name}' created and persisted.")
    else:
        print(f"Attempting to load existing vector store for '{category_name}' from {category_db_path}...")
        if category_db_path.exists() and any(category_db_path.iterdir()):
            vectorstore = Chroma(
                persist_directory=str(category_db_path),
                embedding_function=embeddings
            )
            print(f"Existing vector store for '{category_name}' loaded successfully.")
        else:
            print(f"No documents provided for '{category_name}' and no existing vector store found or it's empty. Skipping.")
            return None
    return vectorstore

def start_llamafile_server():
    """
    Starts the llamafile server as a subprocess and waits for it to become ready.
    Implements a robust retry mechanism, logging llamafile's output to a file.
    """
    llamafile_path = Path(LLAMAFILE_NAME)
    if not llamafile_path.exists():
        print(f"Error: Llamafile '{LLAMAFILE_NAME}' not found in the current directory.")
        print("Please download it from Hugging Face and place it here, then run 'chmod +x {LLAMAFILE_NAME}'.")
        return None, "Llamafile executable not found."

    try:
        from openai import OpenAI
        client = OpenAI(base_url=LLAMAFILE_API_BASE, api_key="sk-no-key-required", timeout=5.0) 
        client.models.list() 
        print("Llamafile server appears to be already running and responsive.")
        return None, None # Return None for process and error if already running
    except Exception:
        pass # Llamafile not running, proceed to start it

    print(f"Attempting to start llamafile server with GPU offload: {LLAMAFILE_NAME} -ngl {GPU_OFFLOAD_LAYERS}...")
    print(f"Llamafile server output will be logged to: {LLAMAFILE_LOG_FILE}")

    log_file_handle = open(LLAMAFILE_LOG_FILE, 'w') # Open in write mode to clear previous logs
    try:
        command_str = ( 
            f"./{llamafile_path.name} --server --host 0.0.0.0 --port 8080 "
            f"-ngl {GPU_OFFLOAD_LAYERS} > {LLAMAFILE_LOG_FILE} 2>&1"
        )
        if GPU_OFFLOAD_LAYERS == 0:
            command_str = f"./{llamafile_path.name} --server --host 0.0.0.0 --port 8080 > {LLAMAFILE_LOG_FILE} 2>&1"

        process = subprocess.Popen(
            command_str, 
            cwd=os.getcwd(),
            shell=True 
        )
        print(f"Llamafile server (with GPU offload) started with PID: {process.pid}.")
        print("Waiting for llamafile server to become ready (up to 240 seconds)...")

        start_time = time.time()
        timeout = 240
        check_interval = 5

        from openai import OpenAI
        while time.time() - start_time < timeout:
            try:
                client = OpenAI(base_url=LLAMAFILE_API_BASE, api_key="sk-no-key-required", timeout=10.0)
                client.models.list() 
                print("\nLlamafile server successfully connected and is ready!")
                return process, None 
            except Exception:
                if process.poll() is not None: 
                    print("\nLlamafile server process terminated unexpectedly during startup.")
                    print(f"Please check '{LLAMAFILE_LOG_FILE}' for error details.")
                    return None, f"Llamafile process terminated unexpectedly. Check {LLAMAFILE_LOG_FILE}"
                time.sleep(check_interval)
        
        if process.poll() is None:
            print("\nLlamafile server did not become ready within the timeout period.")
            print(f"Please check '{LLAMAFILE_LOG_FILE}' for potential startup logs.")
            return None, f"Llamafile server timed out during startup. Check {LLAMAFILE_LOG_FILE}"
        else: 
            print("\nLlamafile server process terminated unexpectedly during startup check (after loop).")
            print(f"Please check '{LLAMAFILE_LOG_FILE}' for error details.")
            return None, f"Llamafile process terminated unexpectedly. Check {LLAMAFILE_LOG_FILE}"

    except Exception as e:
        print(f"Failed to launch llamafile server. Error details: {e}")
        print(f"Please ensure '{LLAMAFILE_NAME}' is executable (chmod +x) and compatible with your system.")
        return None, f"Failed to launch llamafile server: {e}"
    finally:
        log_file_handle.close() 

def initialize_llm_for_summarization():
    """Initializes and returns an LLM instance for summarization within the scraper."""
    try:
        llm = ChatOpenAI(
            model_name="llamafile",
            openai_api_key="sk-no-key-required",
            openai_api_base=LLAMAFILE_API_BASE,
            temperature=0.0, # Keep low for concise summaries
            request_timeout=60.0 # Timeout for LLM calls during scraping
        )
        # Test connection again to be sure after server start
        llm.invoke("Hello", max_tokens=10) 
        print("LLM for scraper summarization initialized successfully.")
        return llm
    except Exception as e:
        print(f"Error initializing LLM for scraper summarization: {e}")
        return None

# --- Helper for Parsing User Input ---
def parse_user_input(user_input, available_kbs_keys):
    """
    Parses user input to determine selected category, query text, and metadata filters.
    Returns (selected_category, original_query_text, query_text_for_llm, metadata_filter).
    """
    selected_category = 'auto' 
    original_query_text = user_input
    query_text_for_llm = user_input 
    metadata_filter = None 
    
    # Check for explicit category prefix
    if ':' in user_input:
        parts = user_input.split(':', 1)
        category_prefix = parts[0].strip().lower()
        if category_prefix in available_kbs_keys:
            selected_category = category_prefix
            original_query_text = parts[1].strip() 
            query_text_for_llm = original_query_text 
        else:
            print(f"{ANSI_BLUE}Invalid category '{category_prefix}'. Attempting intelligent routing.{ANSI_RESET}")
            # selected_category remains 'auto' for intelligent routing/fallback

    # Determine if it's a summary request (news, blog, social media) for auto-routing
    is_summary_req_auto_detected = False
    if selected_category == 'auto' and is_headline_query(original_query_text):
        is_summary_req_auto_detected = True
        # If it's a headline query and no explicit KB was explicitly chosen,
        # then route to direct file system retrieval.
        if not any(prefix.lower() + ':' in original_query_text.lower() for prefix in available_kbs_keys):
            selected_category = 'direct_headlines'
            # Parse date and site filters from the original query, these will be used by get_headlines_from_filesystem
            parsed_filter, cleaned_query_after_filter = parse_date_and_site_from_query(original_query_text)
            if parsed_filter:
                metadata_filter = parsed_filter
                query_text_for_llm = cleaned_query_after_filter # Cleaned query for potential future semantic search if needed
            # If no explicit filters, query_text_for_llm remains original_query_text
        else: # It's a headline query, but an explicit KB was given (e.g., NEWS_SUMMARIES:), so route to that KB's summary handler
            if "blog" in original_query_text.lower():
                selected_category = 'blog_summaries'
            elif "social media" in original_query.lower() or "trending" in original_query_text.lower():
                selected_category = 'social_media_summaries'
            else:
                selected_category = 'news_summaries'
            # Parse date and site filters for the chosen summary KB
            parsed_filter, cleaned_query_after_filter = parse_date_and_site_from_query(original_query_text)
            if parsed_filter:
                metadata_filter = parsed_filter
                query_text_for_llm = cleaned_query_after_filter
            # If no specific filter was found, query_text_for_llm remains original_query_text


    # If an explicit summary KB was selected (e.g., NEWS_SUMMARIES:), apply filters for that KB
    # This block is now redundant because the logic above handles it.
    # if selected_category in ['news_summaries', 'blog_summaries', 'social_media_summaries'] and not is_summary_req_auto_detected:
    #     parsed_filter, cleaned_query_after_filter = parse_date_and_site_from_query(original_query_text)
    #     if parsed_filter:
    #         metadata_filter = parsed_filter
    #         query_text_for_llm = cleaned_query_after_filter 
    #     # If no specific filter was found, query_text_for_llm remains original_query_text

    return selected_category, original_query_text, query_text_for_llm, metadata_filter

# --- New Function: get_headlines_from_filesystem ---
def get_headlines_from_filesystem(original_query_text, metadata_filter):
    """
    Retrieves and returns headlines directly from news JSON files on the file system,
    applying date and site filters. No LLM or ChromaDB involvement.
    """
    print(f"Retrieving headlines directly from file system for '{original_query_text}'...")
    
    retrieved_headlines = []
    
    # Iterate through NEWS_DIR to find timestamped site folders
    dirs_to_walk = []
    potential_subdirs = [d for d in NEWS_DIR.iterdir() if d.is_dir()] 
    if potential_subdirs:
        dirs_to_walk.extend(potential_subdirs)
    else:
        dirs_to_walk = [NEWS_DIR] # Fallback if no timestamped subdirs

    for start_dir in dirs_to_walk:
        for root, _, files in os.walk(start_dir):
            for file_name in files:
                if file_name.endswith(".json"): # Only process JSON metadata files
                    file_path = os.path.join(root, file_name)
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            metadata = json.load(f)
                            
                            # Normalize site_origin from www.domain.com to domain.com if present
                            if "site_origin" in metadata and metadata["site_origin"].startswith("www."):
                                metadata["site_origin"] = metadata["site_origin"][4:]

                            # Apply filters
                            match = True
                            if metadata_filter:
                                for filter_key, filter_value in metadata_filter.items():
                                    if filter_key not in metadata:
                                        match = False
                                        break
                                    
                                    doc_value = metadata[filter_key]
                                    
                                    if filter_key == "publication_date":
                                        # Parse doc_value (e.g., "2025-08-20_14-40-25") into datetime object
                                        try:
                                            # Corrected parsing format for publication_date from JSON
                                            doc_datetime = datetime.strptime(doc_value, "%Y-%m-%d_%H-%M-%S")
                                        except ValueError:
                                            print(f"  Warning: Could not parse publication_date '{doc_value}' from {file_name}. Skipping date filter for this doc.")
                                            match = False # Treat as non-match if date unparseable
                                            break

                                        # Parse filter_value's $gte and $lte into datetime objects
                                        # These filter values come from parse_date_and_site_from_query and are already ISO format
                                        filter_start_dt = datetime.fromisoformat(filter_value["$gte"])
                                        filter_end_dt = datetime.fromisoformat(filter_value["$lte"])

                                        if not (filter_start_dt <= doc_datetime <= filter_end_dt):
                                            match = False
                                            break
                                    elif doc_value != filter_value: # Simple exact match for site_origin
                                        match = False
                                        break
                            
                            if match:
                                title = metadata.get("original_title", "No Title Available")
                                summary = metadata.get("article_summary", "No Summary Available")
                                site_origin = metadata.get("site_origin", "Unknown Site")
                                source_filename = Path(file_path).name # Just the filename for source
                                
                                retrieved_headlines.append({
                                    "title": title,
                                    "summary": summary,
                                    "site_origin": site_origin,
                                    "source": source_filename
                                })
                    except json.JSONDecodeError as e:
                        print(f"  Error reading JSON metadata for {file_name}: {e}. Skipping.")
                    except Exception as e:
                        print(f"  Error processing {file_name}: {e}.")

    # Return the list of headlines instead of printing directly
    return retrieved_headlines


# --- Renamed and Updated Summary Handler Function ---

def handle_summary_query(llm, summary_vectorstore, original_query_text, query_text_for_llm, metadata_filter, HEADLINE_SUMMARIZER_PROMPT, SYSTEM_INSTRUCTION_PROMPT, kb_name_for_print="Summaries"):
    """
    Handles queries specifically for summaries (news, blog, social media).
    This function uses ChromaDB and the LLM for semantic search and synthesis.
    """
    print(f"Searching and generating high-level overview from '{kb_name_for_print.upper()}' knowledge base (streaming)...")
    
    if summary_vectorstore:
        embeddings_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
        
        chroma_client_collection = summary_vectorstore._collection

        chroma_where_clause = {}
        if metadata_filter:
            conditions = []
            for k, v in metadata_filter.items():
                if k == "publication_date" and isinstance(v, dict):
                    # For publication_date, break it into two separate $gte and $lte conditions
                    if "$gte" in v:
                        conditions.append({"publication_date": {"$gte": v["$gte"]}})
                    if "$lte" in v:
                        conditions.append({"publication_date": {"$lte": v["$lte"]}})
                else:
                    conditions.append({k: v})
            
            if len(conditions) == 1:
                chroma_where_clause = conditions[0]
            elif len(conditions) > 1:
                chroma_where_clause = {"$and": conditions}

        print(f"  DEBUG: Initial metadata_filter: {metadata_filter}")
        print(f"  DEBUG: Constructed chroma_where_clause: {chroma_where_clause}")

        retrieved_docs_raw = []
        try:
            final_query_text_for_embedding = query_text_for_llm if query_text_for_llm.strip() else "general news summaries and headlines"
            
            query_embedding = embeddings_model.embed_query(final_query_text_for_embedding)
            
            print(f"  Using ChromaDB .query() for semantic search with filters (if any). Query embedding based on: '{final_query_text_for_embedding}'")
            
            results = chroma_client_collection.query(
                query_embeddings=[query_embedding],
                n_results=100, # Retrieve a large number to ensure filter covers all matches
                where=chroma_where_clause, # Apply the correctly formatted filter
                include=['documents', 'metadatas', 'distances']
            )

            if results['documents']:
                for i in range(len(results['documents'])):
                    doc_content = results['documents'][i][0]
                    doc_metadata = results['metadatas'][i][0]
                    retrieved_docs_raw.append(LcDocument(page_content=doc_content, metadata=doc_metadata))
            
            if not retrieved_docs_raw and metadata_filter:
                raise ValueError(f"ChromaDB .query() returned no results after applying filter {metadata_filter}. This means no documents match these criteria.")
            elif not retrieved_docs_raw:
                raise ValueError("ChromaDB .query() returned no results for semantic search (without explicit filters).")

        except Exception as e:
            print(f"Error during ChromaDB retrieval with .query(): {e}. This might mean no results matching criteria were found or an unexpected error occurred. No fallback to LangChain retriever in this case as it won't magically find documents not indexed.")
            retrieved_docs_raw = []

        # Deduplicate retrieved summaries based on file_path (most reliable)
        unique_docs_map = {}
        for doc in retrieved_docs_raw:
            doc_key = doc.metadata.get("file_path") 
            if doc_key and doc_key not in unique_docs_map: 
                unique_docs_map[doc_key] = doc
            elif not doc_key: 
                content_key = f"{doc.metadata.get('original_title', '')}_{doc.page_content[:50]}"
                if content_key not in unique_docs_map:
                    unique_docs_map[content_key] = doc

        retrieved_docs = list(unique_docs_map.values())
        print(f"  Retrieved {len(retrieved_docs_raw)} raw summaries, kept {len(retrieved_docs)} unique summaries.")

        if not retrieved_docs:
            print(f"{ANSI_BLUE}No relevant summaries found in the '{kb_name_for_print}' knowledge base for your query.{ANSI_RESET}") 
            user_query_content = f"Question: {original_query_text}\nAnswer:"
            messages_for_llm = [
                {"role": "system", "content": SYSTEM_INSTRUCTION_PROMPT},
                {"role": "user", "content": user_query_content}
            ]
        else:
            headline_context_parts = []
            print("\n--- DEBUG: Details of Retrieved Documents (from handle_summary_query, after filter) ---")
            for i, doc in enumerate(retrieved_docs):
                print(f"  Document {i+1}:")
                print(f"    Chroma Doc ID: {doc.metadata.get('id', 'N/A')}") 
                print(f"    File Path: {doc.metadata.get('file_path', 'N/A')}")
                print(f"    Site Origin (from metadata): {doc.metadata.get('site_origin', 'N/A')}")
                print(f"    Original Title (from metadata): {doc.metadata.get('original_title', 'N/A')}")
                print(f"    Full Article Summary (page_content):\n{doc.page_content}\n")
                print(f"    All Metadata: {doc.metadata}")
                print("-" * 20)
            print("--- END DEBUG ---\n")
            
            for doc in retrieved_docs:
                title = doc.metadata.get("original_title", "No Title")
                summary = doc.page_content 
                source_info = doc.metadata.get("source", "Unknown Source")
                site_origin = doc.metadata.get("site_origin", "Unknown Site")
                headline_context_parts.append(f"- Source: {source_info}, Site: {site_origin}, {ANSI_YELLOW}Title:{ANSI_RESET} {ANSI_ORANGE}{title}{ANSI_RESET}, {ANSI_YELLOW}Summary:{ANSI_RESET} {ANSI_LIGHT_BLUE}{summary}{ANSI_RESET}")
            
            context_str = "\n".join(headline_context_parts)
            user_query_content = HEADLINE_SUMMARIZER_PROMPT.format(context=context_str)
            messages_for_llm = [
                {"role": "system", "content": SYSTEM_INSTRUCTION_PROMPT},
                {"role": "user", "content": user_query_content}
            ]
    else:
        print(f"{ANSI_BLUE}{kb_name_for_print} knowledge base not loaded. Cannot provide overview.{ANSI_RESET}") 
        return

    print("\n--- Answer (Streaming) ---")
    print(ANSI_BLUE, end="", flush=True) 
    full_response_content = ""
    for chunk in llm.stream(messages_for_llm):
        if chunk.content:
            print(chunk.content, end="", flush=True)
            full_response_content += chunk.content
    print(ANSI_RESET) 
    print("\n")

def handle_general_query(llm, knowledge_bases_dict, original_query_text, query_text_for_llm, target_kb_names, SYSTEM_INSTRUCTION_PROMPT):
    """
    Handles general document queries, which might target a specific KB or all content KBs.
    `target_kb_names` can be a single KB name string or a list of KB names.
    """
    retrieved_docs = []
    
    if isinstance(target_kb_names, str): # Single KB specified
        kb_name = target_kb_names
        current_vectorstore = knowledge_bases_dict.get(kb_name)
        if current_vectorstore:
            print(f"Searching and generating response from '{kb_name.upper()}' knowledge base (streaming)...")
            retriever = current_vectorstore.as_retriever()
            retrieved_docs.extend(retriever.get_relevant_documents(query_text_for_llm))
        else:
            print(f"{ANSI_BLUE}Knowledge base '{kb_name}' not loaded. Skipping search.{ANSI_RESET}")
    elif isinstance(target_kb_names, list): # Multiple KBs (e.g., for 'all' or 'auto' fallback)
        print(f"Searching across {ANSI_BLUE}multiple knowledge bases:{', '.join(kb.upper() for kb in target_kb_names)}{ANSI_RESET} (streaming)...")
        for kb_name in target_kb_names:
            current_vectorstore = knowledge_bases_dict.get(kb_name)
            if current_vectorstore:
                print(f"  Querying {kb_name.upper()} knowledge base...")
                retriever = current_vectorstore.as_retriever()
                retrieved_docs.extend(retriever.get_relevant_documents(query_text_for_llm))
            else:
                print(f"  {ANSI_BLUE}Knowledge base '{kb_name}' not loaded. Skipping.{ANSI_RESET}")

    if not retrieved_docs:
        print(f"{ANSI_BLUE}No relevant documents found in the selected knowledge base(s) for your query.{ANSI_RESET}") 
        context_str = "" # Explicitly empty context for LLM to use general knowledge
        user_query_content = f"Question: {original_query_text}\nAnswer:"
        messages_for_llm = [
            {"role": "system", "content": SYSTEM_INSTRUCTION_PROMPT},
            {"role": "user", "content": user_query_content}
        ]
    else:
        context_parts = []
        for doc in retrieved_docs:
            source_info = doc.metadata.get("source", "Unknown Source")
            original_title = doc.metadata.get("original_title")
            title_part = f", Title: {original_title}" if original_title else ""
            context_parts.append(f"Source: {source_info}{title_part}\nContent: {doc.page_content}")
        context_str = "\n\n".join(context_parts)
        
        user_query_content = f"""Context:
{context_str}

Question: {original_query_text} 

Please summarize the information in the provided context, identifying and creating distinct, relevant sections based on the content. Provide key points under each section in bullet points. If information for a category is not found, you can omit that category or state that the information is not available in the document.

Answer:"""
        messages_for_llm = [
            {"role": "system", "content": SYSTEM_INSTRUCTION_PROMPT},
            {"role": "user", "content": user_query_content}
        ]

    print("\n--- Answer (Streaming) ---")
    print(ANSI_BLUE, end="", flush=True) 
    full_response_content = ""
    for chunk in llm.stream(messages_for_llm):
        if chunk.content:
            print(chunk.content, end="", flush=True)
            full_response_content += chunk.content
    print(ANSI_RESET) 
    print("\n")

# New debug function to inspect ChromaDB content directly
def debug_chroma_content(vectorstore, filter_site_origin):
    """
    Directly queries a ChromaDB vectorstore with a site_origin filter
    and prints the details of the retrieved documents.
    """
    if not vectorstore:
        print(f"Debug: Vectorstore is not initialized. Cannot debug content for {filter_site_origin}.")
        return

    print(f"\n--- DEBUG: Directly Querying ChromaDB for site_origin: {filter_site_origin} ---")
    
    try:
        chroma_client_collection = vectorstore._collection

        results = chroma_client_collection.get(
            where={"site_origin": filter_site_origin},
            limit=100 
        )

        if not results['ids']:
            print(f"  No documents found in ChromaDB with site_origin '{filter_site_origin}'.")
        else:
            print(f"  Found {len(results['ids'])} documents matching site_origin '{filter_site_origin}':")
            for i in range(len(results['ids'])):
                doc_id = results['ids'][i]
                metadata = results['metadatas'][i]
                document_content = results['documents'][i] 

                print(f"  Document {i+1} (Chroma Internal ID: {doc_id}):")
                print(f"    Metadata: {metadata}")
                print(f"    Content (full):\n{document_content}\n") # Print full content for debug
                print("-" * 20)
    except Exception as e:
        print(f"  Error during direct ChromaDB debug query: {e}")

    print("--- END DIRECT CHROMA DEBUG ---\n")


def main():
    """
    Main function to run the RAG application.
    """
    print("✨ Starting Offline RAG Summarizer (Llamafile Version) ✨")

    llamafile_process, llama_startup_error = start_llamafile_server()
    if llama_startup_error:
        print(f"Fatal error during Llamafile server startup: {llama_startup_error}")
        return

    BASE_DATA_DIR.mkdir(exist_ok=True)
    MISC_DIR.mkdir(exist_ok=True)
    SEC_DIR.mkdir(exist_ok=True)
    LEGAL_DIR.mkdir(exist_ok=True)
    NEWS_DIR.mkdir(exist_ok=True)
    BLOG_DIR.mkdir(exist_ok=True)
    SOCIAL_MEDIA_DIR.mkdir(exist_ok=True)
    
    BASE_VECTOR_DB_DIR.mkdir(exist_ok=True)

    NEWS_SUMMARIES_KB_NAME = "news_summaries" 
    BLOG_SUMMARIES_KB_NAME = "blog_summaries"
    SOCIAL_MEDIA_SUMMARIES_KB_NAME = "social_media_summaries"

    print("\n--- Document Organization ---")
    print(f"  Place miscellaneous documents (TXT, DOCX, PDF, generic HTML) in: {MISC_DIR}") 
    print(f"  Place SEC filings (HTML/HTM) in: {SEC_DIR}")
    print(f"  Place legal documents (TXT, DOCX, PDF) in: {LEGAL_DIR}")
    print(f"  Place news articles (TXT, from scraper with JSON metadata including summaries) in: {NEWS_DIR}") 
    print(f"  NEW: Place blog articles (TXT, with JSON metadata including summaries) in: {BLOG_DIR}")
    print(f"  NEW: Place social media content (TXT, with JSON metadata including summaries) in: {SOCIAL_MEDIA_DIR}")
    print("-----------------------------\n")

    knowledge_bases = {}
    
    print(f"Loading documents and initializing vector store for '{MISC_DIR.name}'...")
    misc_docs = load_documents_from_directory(MISC_DIR)
    misc_vectorstore = initialize_vector_store_for_category(misc_docs, MISC_DIR.name)
    if misc_vectorstore:
        knowledge_bases['misc'] = misc_vectorstore
    
    print(f"\nLoading documents and initializing vector store for '{SEC_DIR.name}'...")
    sec_docs = load_documents_from_directory(SEC_DIR)
    sec_vectorstore = initialize_vector_store_for_category(sec_docs, SEC_DIR.name)
    if sec_vectorstore:
        knowledge_bases['sec'] = sec_vectorstore

    print(f"\nLoading documents and initializing vector store for '{LEGAL_DIR.name}'...")
    legal_docs = load_documents_from_directory(LEGAL_DIR)
    legal_vectorstore = initialize_vector_store_for_category(legal_docs, LEGAL_DIR.name)
    if legal_vectorstore:
        knowledge_bases['legal'] = legal_vectorstore

    print(f"\nLoading documents and initializing vector store for '{NEWS_DIR.name}' (full articles)...")
    news_full_docs = load_documents_from_directory(NEWS_DIR, load_summaries=False)
    news_full_vectorstore = initialize_vector_store_for_category(news_full_docs, NEWS_DIR.name, use_summaries=False)
    if news_full_vectorstore:
        knowledge_bases['news_full'] = news_full_vectorstore

    print(f"\nLoading documents and initializing vector store for '{NEWS_SUMMARIES_KB_NAME}' (summaries)...")
    news_summary_docs = load_documents_from_directory(NEWS_DIR, load_summaries=True, summary_kb_name=NEWS_SUMMARIES_KB_NAME) 
    news_summary_vectorstore = initialize_vector_store_for_category(news_summary_docs, NEWS_SUMMARIES_KB_NAME, use_summaries=True)
    if news_summary_vectorstore:
        knowledge_bases[NEWS_SUMMARIES_KB_NAME] = news_summary_vectorstore

    print(f"\nLoading documents and initializing vector store for '{BLOG_SUMMARIES_KB_NAME}' (summaries)...")
    blog_summary_docs = load_documents_from_directory(BLOG_DIR, load_summaries=True, summary_kb_name=BLOG_SUMMARIES_KB_NAME) 
    blog_summary_vectorstore = initialize_vector_store_for_category(blog_summary_docs, BLOG_SUMMARIES_KB_NAME, use_summaries=True)
    if blog_summary_vectorstore:
        knowledge_bases[BLOG_SUMMARIES_KB_NAME] = blog_summary_vectorstore

    print(f"\nLoading documents and initializing vector store for '{SOCIAL_MEDIA_SUMMARIES_KB_NAME}' (summaries)...")
    social_media_summary_docs = load_documents_from_directory(SOCIAL_MEDIA_DIR, load_summaries=True, summary_kb_name=SOCIAL_MEDIA_SUMMARIES_KB_NAME) 
    social_media_summary_vectorstore = initialize_vector_store_for_category(social_media_summary_docs, SOCIAL_MEDIA_SUMMARIES_KB_NAME, use_summaries=True)
    if social_media_summary_vectorstore:
        knowledge_bases[SOCIAL_MEDIA_SUMMARIES_KB_NAME] = social_media_summary_vectorstore


    if not knowledge_bases:
        print("\nNo active knowledge bases found. Please ensure documents are in the correct 'data' subfolders and try again.")
        if llamafile_process and isinstance(llamafile_process, subprocess.Popen):
            llamafile_process.terminate()
        return

    print(f"\nSuccessfully loaded {len(knowledge_bases)} knowledge bases: {', '.join(knowledge_bases.keys())}")

    if 'news_summaries' in knowledge_bases:
        debug_chroma_content(knowledge_bases['news_summaries'], 'cnn.com')
        debug_chroma_content(knowledge_bases['news_summaries'], 'foxnews.com')


    print(f"Connecting to LLM via llamafile at {LLAMAFILE_API_BASE}...")
    try:
        llm = ChatOpenAI(
            model_name="llamafile",
            openai_api_key="sk-no-key-required",
            openai_api_base=LLAMAFILE_API_BASE,
            temperature=0.0, # Keep low for concise summaries
            streaming=True,
            request_timeout=120.0
        )
        print("LLM (via llamafile) initialized successfully!")
    except Exception as e:
        print(f"Error connecting to llamafile LLM at {LLAMAFILE_API_BASE}.")
        print(f"Error details: {e}")
        print(f"Please ensure '{LLAMAFILE_NAME}' is running in server mode and reachable at {LLAMAFILE_API_BASE}.")
        if llamafile_process and isinstance(llamafile_process, subprocess.Popen):
            llamafile_process.terminate()
        return

    print("\nReady to answer questions about your documents!")
    print("Available knowledge bases:")
    for key in knowledge_bases.keys():
        print(f"  - {key.upper()}")
    print("Type 'exit' or 'quit' to end the session.")
    print("To query a specific knowledge base for *detailed* info, type '<CATEGORY>: <your question>'. E.g., 'NEWS_FULL: Tell me about the energy bill.'")
    print("For a high-level summary overview (news, blogs, social media), just type 'Summarize today's headlines' or 'Summarize recent blogs' (no category prefix needed, or use 'NEWS_SUMMARIES:', 'BLOG_SUMMARIES:', etc.)")
    print("If no category is specified and it's not a summary query, I will query across **ALL** content-based knowledge bases (MISC, SEC, LEGAL, NEWS_FULL).") 
    print(f"To debug ChromaDB content directly, type '{ANSI_YELLOW}debug: <site_origin>'{ANSI_RESET} (e.g., 'debug: foxnews.com').")
    print(f"Press {ANSI_YELLOW}Ctrl+C{ANSI_RESET} to stop a streamed response early.") 

    try:
        while True:
            user_input = input(f"\n{ANSI_YELLOW}Your question (e.g., NEWS: Summarize politics): {ANSI_RESET}").strip() 
            if user_input.lower() in ["exit", "quit"]:
                print("Exiting RAG Summarizer. Goodbye!")
                break
            if not user_input:
                print("Please enter a question.")
                continue
            
            if user_input.lower() == "help":
                print("\n--- Help: Available Knowledge Bases & Sample Prompts ---")
                print(f"Available Knowledge Bases (Categories): {', '.join(kb.upper() for kb in knowledge_bases.keys())}")
                print("\nSample Prompts:")
                print(f"{ANSI_YELLOW}General Summary Overview (News, Blogs, Social Media):{ANSI_RESET}") 
                print("  - Summarize today's headlines")
                print("  - What's the news from last week from CNN?")
                print("  - Give me headlines about politics from yesterday.")
                print("  - NEWS_SUMMARIES: Recent economic news.")
                print("  - Summarize recent blogs.")
                print("  - BLOG_SUMMARIES: What are the latest tech blog posts?")
                print("  - Summarize social media trends from today.")
                print("  - SOCIAL_MEDIA_SUMMARIES: What are people saying about stock X?")
                print(f"{ANSI_YELLOW}Detailed Document Query (Full Content):{ANSI_RESET}") 
                print("  - NEWS_FULL: What are the details about the latest energy bill?")
                print("  - NEWS_FULL: Explain the recent market fluctuations from foxnews.com.")
                print("  - MISC: What is the main topic of the document 'my_research_paper.txt'?")
                print("  - SEC: Summarize the latest 10-K filing from Apple Inc.")
                print("  - LEGAL: Explain the concept of 'force majeure' from my legal documents.")
                print(f"\n{ANSI_YELLOW}Tips:{ANSI_RESET}") 
                print("  - Use `CATEGORY:` prefix for specific knowledge base queries.")
                print("  - For dates, you can use 'today', 'yesterday', 'last week', 'this month', 'YYYY-MM-DD', or 'MM-DD-YYYY'.")
                print("  - For sites, you can use 'cnn.com', 'foxnews.com', 'reddit.com', 'medium.com', etc.")
                print("-----------------------------------------------------")
                continue

            if user_input.lower().startswith("debug:"):
                site_to_debug = user_input[len("debug:"):].strip()
                if site_to_debug and 'news_summaries' in knowledge_bases:
                    debug_chroma_content(knowledge_bases['news_summaries'], site_to_debug)
                else:
                    print(f"{ANSI_BLUE}Invalid debug command or 'news_summaries' KB not loaded. Usage: debug: <site_origin>{ANSI_RESET}")
                continue

            selected_category, original_query_text, query_text_for_llm, metadata_filter = \
                parse_user_input(user_input, knowledge_bases.keys())

            # NEW: Handle 'direct_headlines' category for file system read
            if selected_category == 'direct_headlines':
                print(f"Retrieving headlines directly from file system for '{original_query_text}'...")
                # The get_headlines_from_filesystem function will handle filtering and returning data
                retrieved_headlines_list = get_headlines_from_filesystem(original_query_text, metadata_filter)
                
                if not retrieved_headlines_list:
                    print(f"{ANSI_BLUE}No headlines found matching your criteria from the file system.{ANSI_RESET}")
                else:
                    # Format the retrieved headlines for the LLM's context
                    llm_context_for_executive_summary = []
                    for item in retrieved_headlines_list:
                        llm_context_for_executive_summary.append(
                            f"- Source: {item['source']}, Site: {item['site_origin']}, Title: {item['title']}, Summary: {item['summary']}"
                        )
                    context_str_for_llm = "\n".join(llm_context_for_executive_summary)

                    # Prepare messages for LLM to generate the executive summary
                    executive_summary_prompt = f"""You are a brilliant news analyst.
Given the following list of news article titles and their brief summaries, provide a concise executive summary of the key themes and most important developments.
Group related articles by theme and provide a brief overview for each theme.
Do NOT list individual articles in this executive summary. Focus on the overarching narrative.

Input Headlines and Summaries:
{context_str_for_llm}

Executive Summary:"""
                    
                    messages_for_llm = [
                        {"role": "system", "content": SYSTEM_INSTRUCTION_PROMPT},
                        {"role": "user", "content": executive_summary_prompt}
                    ]

                    # Print the LLM's executive summary first
                    print("\n--- Executive News Briefing ---")
                    print(ANSI_BLUE, end="", flush=True)
                    for chunk in llm.stream(messages_for_llm):
                        if chunk.content:
                            print(chunk.content, end="", flush=True)
                    print(ANSI_RESET)
                    print("\n--- End Executive Briefing ---\n")

                    # Then, print the original list of headlines
                    print("\n--- Retrieved Headlines (Direct from File System) ---")
                    for i, item in enumerate(retrieved_headlines_list):
                        print(f"{i+1}. {ANSI_YELLOW}Title:{ANSI_RESET} {ANSI_ORANGE}{item['title']}{ANSI_RESET}")
                        print(f"   {ANSI_YELLOW}Summary:{ANSI_RESET} {ANSI_LIGHT_BLUE}{item['summary']}{ANSI_RESET}")
                        print(f"   (Source: {item['source']}, Site: {item['site_origin']})")
                        print("-" * 20)
                    print("--- END HEADLINES ---")
                
                continue # Skip the rest of the loop and prompt for next question

            # Existing logic for summary KBs (now only for LLM synthesis)
            elif selected_category in [NEWS_SUMMARIES_KB_NAME, BLOG_SUMMARIES_KB_NAME, SOCIAL_MEDIA_SUMMARIES_KB_NAME]:
                if selected_category in knowledge_bases:
                    handle_summary_query(llm, knowledge_bases[selected_category], 
                                         original_query_text, query_text_for_llm, 
                                         metadata_filter, HEADLINE_SUMMARIZER_PROMPT, 
                                         SYSTEM_INSTRUCTION_PROMPT, selected_category)
                else:
                    print(f"{ANSI_BLUE}{selected_category.replace('_', ' ').title()} knowledge base not loaded. Cannot provide overview.{ANSI_RESET}")
            # If not a summary query, proceed to general content query logic
            elif selected_category in ['misc', 'sec', 'legal', 'news_full']:
                handle_general_query(llm, knowledge_bases, original_query_text, 
                                     query_text_for_llm, selected_category, SYSTEM_INSTRUCTION_PROMPT)
            elif selected_category == 'auto':
                content_kbs_for_auto_query = [
                    kb for kb in knowledge_bases if kb not in [NEWS_SUMMARIES_KB_NAME, BLOG_SUMMARIES_KB_NAME, SOCIAL_MEDIA_SUMMARIES_KB_NAME]
                ]
                if content_kbs_for_auto_query:
                    handle_general_query(llm, knowledge_bases, original_query_text, 
                                         query_text_for_llm, content_kbs_for_auto_query, 
                                         SYSTEM_INSTRUCTION_PROMPT)
                else:
                    print(f"{ANSI_BLUE}No general content knowledge bases loaded to answer your query.{ANSI_RESET}")
            else:
                print(f"{ANSI_BLUE}Could not determine query type or find relevant knowledge base for '{user_input}'. Please try 'help' for options.{ANSI_RESET}")

    except KeyboardInterrupt:
        print(f"{ANSI_RESET}\nResponse generation stopped by user.{ANSI_RESET}") 
    except Exception as e:
        print(f"{ANSI_RESET}An error occurred during processing: {e}{ANSI_RESET}") 
        print(f"{ANSI_RESET}Please ensure the llamafile server is still running and try again.{ANSI_RESET}") 
    finally:
        if llamafile_process and isinstance(llamafile_process, subprocess.Popen):
            print("Terminating llamafile server...")
            llamafile_process.terminate()
            llamafile_process.wait(timeout=5)
            if llamafile_process.poll() is None:
                print("Llamafile process did not terminate gracefully, forcing kill.")
                llamafile_process.kill()
        else:
            print("Llamafile process was not running or not successfully started by app.py. No process to terminate.")

if __name__ == "__main__":
    main()

